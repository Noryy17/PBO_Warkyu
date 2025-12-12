#!/usr/bin/env python3
"""
Convert LAPORAN_PERANCANGAN_SISTEM.md to DOCX format
Following template formatting from ContohTemplateMakalah.docx
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Create border elements
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, style='Normal', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                            bold=False, font_size=12, font_name='Times New Roman'):
    """Add a paragraph with specific formatting"""
    para = doc.add_paragraph(text, style=style)
    para.alignment = alignment
    
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        # Set font for Asian text
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rFonts)
    
    return para


def add_image_if_exists(doc, image_path, width_cm=14):
    """Add image to document if it exists"""
    if os.path.exists(image_path):
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(image_path, width=Cm(width_cm))
            return True
        except Exception as e:
            print(f"Warning: Could not add image {image_path}: {e}")
            return False
    else:
        print(f"Warning: Image not found: {image_path}")
        return False


def parse_markdown_table(lines):
    """Parse markdown table to extract rows"""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # Skip separator lines
            if re.match(r'^\|[\s\-:|]+\|$', line):
                continue
            # Extract cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_table_from_markdown(doc, table_lines):
    """Add a table from markdown format"""
    rows_data = parse_markdown_table(table_lines)
    if not rows_data:
        return
    
    # Create table
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Table Grid'
    
    # Fill table
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # Format cell text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    # Bold for header row
                    if i == 0:
                        run.font.bold = True


def convert_markdown_to_docx(md_file, output_file, base_dir):
    """Main conversion function"""
    
    # Create document
    doc = Document()
    
    # Set margins (2.54 cm = 1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # State variables
    in_table = False
    table_lines = []
    in_code_block = False
    bab_counter = 1
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            i += 1
            continue
        
        # Handle tables
        if '|' in line and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue
        
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                add_table_from_markdown(doc, table_lines)
                doc.add_paragraph()  # Add spacing after table
                in_table = False
                table_lines = []
                continue
        
        # Handle main title (first # heading)
        if i < 5 and line.startswith('# ') and 'LAPORAN' in line.upper():
            add_formatted_paragraph(doc, line[2:].strip(), alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   bold=True, font_size=14)
            i += 1
            continue
        
        # Handle subtitle (second # heading - Sistem Digital Wakrun)
        if i < 5 and line.startswith('# ') and 'SISTEM' in line.upper():
            add_formatted_paragraph(doc, line[2:].strip(), alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   bold=True, font_size=13)
            doc.add_paragraph()  # Add spacing
            i += 1
            continue
        
        # Handle metadata (Mata Kuliah, Sistem, Tanggal, Versi)
        if line.startswith('**') and ':' in line and i < 15:
            # Extract bold label and value
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   font_size=11)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['---', '___', '***']:
            doc.add_paragraph()  # Just add spacing
            i += 1
            continue
        
        # Handle major sections (## 1. PENDAHULUAN, etc.)
        if re.match(r'^## \d+\.', line):
            doc.add_paragraph()  # Add spacing before
            section_title = re.sub(r'^## \d+\.\s*', '', line)
            
            # Convert to BAB format
            bab_text = f"BAB {['I', 'II', 'III', 'IV', 'V', 'VI', 'VII'][bab_counter-1] if bab_counter <= 7 else str(bab_counter)}"
            add_formatted_paragraph(doc, bab_text, alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   bold=True, font_size=12)
            add_formatted_paragraph(doc, section_title.upper(), alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                   bold=True, font_size=12)
            doc.add_paragraph()
            bab_counter += 1
            i += 1
            continue
        
        # Handle subsections (### 1.1, ### 3.1, etc.)
        if line.startswith('### '):
            subsection_title = line[4:].strip()
            add_formatted_paragraph(doc, subsection_title, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                                   bold=True, font_size=12)
            i += 1
            continue
        
        # Handle sub-subsections (#### or #####)
        if line.startswith('#### ') or line.startswith('##### '):
            level = 4 if line.startswith('#### ') else 5
            subsection_title = line[level+1:].strip()
            para = add_formatted_paragraph(doc, subsection_title, 
                                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                                          bold=True, font_size=11 if level == 4 else 11)
            # Add some spacing for sub-subsections
            if level == 5:
                para.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue
        
        # Handle images
        if line.startswith('![') and '](' in line:
            # Extract image path
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption = match.group(1)
                img_path = match.group(2)
                full_img_path = os.path.join(base_dir, img_path)
                
                # Add image
                doc.add_paragraph()  # Spacing before
                if add_image_if_exists(doc, full_img_path, width_cm=13):
                    # Add caption
                    caption_para = add_formatted_paragraph(doc, f"Gambar: {caption}", 
                                                          alignment=WD_ALIGN_PARAGRAPH.CENTER, 
                                                          font_size=10)
                    caption_para.runs[0].font.italic = True
                doc.add_paragraph()  # Spacing after
            i += 1
            continue
        
        # Handle lists (bullet points with - or numbered)
        if re.match(r'^\s*[-*]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            # Extract indent level and text
            indent_match = re.match(r'^(\s*)', line)
            indent_level = len(indent_match.group(1)) // 2 if indent_match else 0
            
            text = re.sub(r'^\s*[-*\d.]+\s+', '', line)
            # Remove markdown bold/italic
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            
            para = add_formatted_paragraph(doc, f"• {text}", 
                                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                                          font_size=12)
            para.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip() and not line.startswith('#'):
            # Remove markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            
            # Check if this is a "Penanggung Jawab" section
            if '**Penanggung Jawab:**' in line or 'Penanggung Jawab:' in text:
                add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                                       bold=True, font_size=11)
            else:
                add_formatted_paragraph(doc, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                                       font_size=12)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)
    
    # Save document
    doc.save(output_file)
    print(f"✓ Document saved to: {output_file}")


if __name__ == '__main__':
    # Paths
    script_dir = Path(__file__).parent
    md_file = script_dir / 'LAPORAN_PERANCANGAN_SISTEM.md'
    output_file = script_dir / 'LAPORAN_PERANCANGAN_SISTEM.docx'
    
    print("Starting conversion...")
    print(f"Input: {md_file}")
    print(f"Output: {output_file}")
    print()
    
    convert_markdown_to_docx(md_file, output_file, script_dir)
    
    # Verify output
    if output_file.exists():
        size_mb = output_file.stat().st_size / (1024 * 1024)
        print(f"✓ File created successfully")
        print(f"✓ File size: {size_mb:.2f} MB")
        
        # Count images
        from docx import Document
        doc = Document(str(output_file))
        img_count = len(doc.inline_shapes)
        print(f"✓ Images embedded: {img_count}")
    else:
        print("✗ Error: Output file was not created")
